#!/usr/bin/python
from django.shortcuts import render, redirect
from django.views.decorators.csrf import csrf_exempt
from django.http import JsonResponse, StreamingHttpResponse
from django.core.serializers import serialize
# 登录检查
from django.contrib.auth.decorators import login_required
# 限制访问方式
from django.views.decorators.http import require_http_methods
# 引入数据库
from . import models
from . import task
# from . import to_pdf
import docx
import random
import subprocess
import re
from . import corpus
import time
from .task import fanyi, enfanyi, zh_sg, en_sg,de_zh_sg,de_en_sg,de_en_fanyi,de_zh_fanyi
from docx import Document
from docx.document import Document as _Document
from docx.oxml.text.paragraph import CT_P
from docx.oxml.table import CT_Tbl
from docx.table import _Cell, Table, _Row
from docx.text.paragraph import Paragraph
from os.path import basename
from django.db.models import Q
import copy
from django.shortcuts import get_object_or_404
import langdetect
# 首页面
@login_required
def index(request):
    files = models.file_information.objects.filter(file_user=request.user)
    return render(request, 'work/index.html', {"files":files})


# 上传首页面
@login_required
def update(request):
    return render(request, 'work/update.html', {})

# ################################### 文件上传 ####################################### #
# 文件基本上传
@login_required
@require_http_methods(["POST"])
def file_update(request):
    # try:
    print("开始获取")
    # 获取文件和名称
    file = request.FILES.get('file')
    name = file.name
    # 获取当前的用户
    user = request.user
    print("获取到，存储")
    print(file, name, user)
    # 存入数据库
    update_file = models.file_information(file_name = name, file=file, file_user=user)
    print("???")
    print(update_file)
    update_file.save()
    # 根据文件位置进行修改
    doc_path = "/home/admin/second_Edition/" + str(update_file.file)
    print(doc_path)
    id = update_file.id
    # 如果上传的文件不是.docx文件 转换为doc文件
    # if doc_path.split('.', -1)[-1] == "doc":
    #     print("不是docx")
    #     a = subprocess.check_output(["soffice", "--headless",
    #                                  "--invisible", "--convert-to",
    #                                  "docx", doc_path,
    #                                  "--outdir", "/home/admin/second_Edition/static/work/doc"])
    #     name = str(update_file.file) + "x"
    #     print("docx" + name)
    #     print(a)
    #     print(name)
    #     docx_save = models.file_information.objects.get(pk=id)
    #     docx_save.file = name
    #     docx_save.save()

    # 检测该word的文本是那种语言
    doc = docx.Document(file)
    paragraphs = doc.paragraphs
    # 从全部段落中随机提取一段
    print(len(paragraphs), '=' * 90)
    try:
        text = paragraphs[random.randint(0,len(paragraphs))]
        print(len(paragraphs), '-' * 90)
        while len(text.text) < 1:
            print(len(paragraphs),'='*90)
            text = paragraphs[random.randint(0, len(paragraphs))]
            if len(text.text) >= 1:
                break
        file_language = langdetect.detect(text.text)
    except:
        file_language = ''

    print(file_language,'*'*90)
    print("end")
    return render(request, 'work/file_info.html', {"msg": file.name, 'id': update_file.id,'file_language':file_language})


# 文件信息填写
@login_required
@require_http_methods(["POST"])
def file_info(request):
    name = request.POST['name']
    info = request.POST['info']
    over_time = request.POST['over_time']
    id = request.POST['id']     # 文件ID
    print(id)
    yuan = request.POST['yuan']
    print("源："+yuan)
    mubiao = request.POST['mubiao']
    print("目标："+ mubiao)
    print("时间", over_time)
    # 查询当前的数据
    file = models.file_information.objects.get(pk=id)
    file.file_name = name
    file.file_info = info
    file.yuan = yuan
    file.mubiao = mubiao
    file.over_time = over_time
    file.save()
    print("保存完成")
    return render(request, 'work/trans_type.html', {'id': id})


# 选择翻译种类
@login_required
@require_http_methods(["POST"])
def trans_type(request):
    id = request.POST['id']
    zl = request.POST['zl']
    print("翻译的种类"+zl)
    file = models.file_information.objects.get(pk=id)
    file.fanyi_zl= zl
    file.save()
    print(file)
    fx = file.mubiao
    yuan = file.yuan
    # fx 1 == 英 -》中
    # fx 2 == 中 -》英
    print("fx", fx)
    if fx == 2 and yuan == 1:
        type = 'work'
    if fx == 1 and yuan == 2:
        type = 'enwork'
    if yuan == 3 and fx == 1:
        type = 'de_zh'
    if yuan == 3 and fx == 2:
        type = 'de_en'
    return render(request, 'work/loading.html', {"file": file, 'type': type})


@login_required
def work(request):
    try:
        id = request.GET['id']
        type = request.GET['type']
        yuanwen = models.YuanWenTable.objects.filter(file_id=id).values('id', 'yuanwen')[:100]

        if len(yuanwen) >= 1:
            yiwen = models.YiWenTable.objects.filter(file=id).values('id', 'yiwen_wait_over')[:100]
            file = models.file_information.objects.get(pk=id)
            print(yuanwen)
            print(yiwen)
            all_wenjian = zip(yuanwen, yiwen)
            count = len(yuanwen)
            print(count, '=' * 70)
            file_id = file.id
            # file.work_count += 1
            # file.save()
            if type == 'work':
                return render(request, 'work/work.html', {"wenjian": all_wenjian, "id": id, "file": file, 'file_id':file_id, 'count': count})
            if type == 'enwork':
                return render(request, 'work/enwork.html', {"wenjian": all_wenjian, "id": id, "file": file, 'file_id':file_id,  'count': count})
            else:
                return render(request, 'work/base_fanyi.html',{"wenjian": all_wenjian, "id": id, "file": file, 'file_id': file_id, 'count': count, 'type':type })

    except Exception as e:

        print(e)
        return render(request, 'work/work_error.html', {})


# ################################### 机器预编译 ####################################### #
# 中-》英 编译等待界面
# @login_required
# def work_wait(request):
#     return render(request, 'work/work_wait.html', {})
def show_wait(request):
    return render(request, 'work/loading.html', {})


def show_new_work(request):
    id = request.GET['id']
    global count
    yuanwen = models.YuanWenTable.objects.filter(file_id=id).values('id', 'yuanwen')[:100]
    if len(yuanwen) > 1:
        yiwen = models.YiWenTable.objects.filter(file=id).values('id', 'yiwen_wait_over')[:100]
        file = models.file_information.objects.get(pk=id)
        wenjian = zip(yuanwen, yiwen)
        count = len(yuanwen)
        print(count,'='*70)
    return render(request, 'work/new_work.html', {"wenjian": wenjian, "id": id,  "file": file,'count':count})


def read_item(parent):
    if isinstance(parent, _Document):
        parent_elm = parent.element.body
    elif isinstance(parent, _Cell):
        parent_elm = parent._tc
    elif isinstance(parent, _Row):
        parent_elm = parent._tr
    else:
        raise ValueError("something's not right")
    for child in parent_elm.iterchildren():
        if isinstance(child, CT_P):
            yield Paragraph(child, parent)
        elif isinstance(child, CT_Tbl):
            yield Table(child, parent)


# 获取当前文章总共有多少段
def get_all_paragraph_count(id):
    # 定义存储所有段落的列表
    paragraph_list = []
    # 查询ID下的文件
    word_file = models.file_information.objects.get(pk=id)
    # 得到文件路径
    path = word_file.file
    # docx操作docx文档
    file = docx.Document(path)
    count_nums = 0
    for block in read_item(file):
        # 读取段落 并进行翻译 和存储
        if isinstance(block, Paragraph):
            count_nums += 1
        # 读取表格和存储单元格数据
        elif isinstance(block, Table):
            count_nums += 1
    word_file.paragraph_count = count_nums
    word_file.save()
    print(word_file.paragraph_count, '/'*9)
    print("文章总长", count_nums,'*'*9)
    return count_nums


def get_all_par(request):
    ret= {'code':0}
    id = request.GET['id']
    # paragraph_count = get_all_paragraph_count(id)
    file_info = models.file_information.objects.get(pk=id)
    paragraph_count = file_info.paragraph_count
    ret['data'] = paragraph_count
    return JsonResponse(ret)


# 顺序读取wordneir
def iter_block_items(parent, file_id, user):
    """
    在*parent*中生成对每个段落和表子级的引用，按文档顺序。
    每个返回值都是表或段落。
    *parent*通常是对主级的引用文档对象，但也适用于单元格对象，它本身可以包含段落和表格。
    """
    if isinstance(parent, _Document):
        parent_elm = parent.element.body
    elif isinstance(parent, _Cell):
        parent_elm = parent._tc
    elif isinstance(parent, _Row):
        parent_elm = parent._tr
    else:
        raise ValueError("something's not right")
    for child in parent_elm.iterchildren():
        if isinstance(child, CT_P):
            yield Paragraph(child, parent)
        elif isinstance(child, CT_Tbl):
            yield Table(child, parent)


# 异步任务：讲文件读取和存储放入异步任务
from . import translationTask
def red_translation(request,id):
    ret = {'code':0}
    # id = request.GET.get('id')
    ret['id'] = id
    print('开始执行异步任务','*'*9)
    user_id = request.user.id
    translationTask.pre_translation.delay(user_id,id)
    return JsonResponse(ret)


page_count = 50
page_start = 0
count = 0
from django.db import connection
def ajax_display(request):

    # 获取前段传过来的查询起始值
    nums = request.GET.get('nums')
    # 获取前端传过来的文件ID
    file_id = request.GET.get('file_id')
    ret = {'code':0}
    n_nums = int(nums)
    # 起始页面等于上次起始页面+每次的页面数
    new_nums = n_nums+page_count
    ret['nums'] = new_nums

    sql1 = 'select id,yuanwen from work_yuanwentable WHERE file_id = %s LIMIT %s, %s'%(file_id,n_nums,page_count)
    sql2 = 'select id,yiwen_wait_over from work_yiwentable WHERE file_id = %s LIMIT %s, %s'%(file_id,n_nums,page_count)
    #
    num_sql = 'select file_name,numing, all_num, paraing, all_para,jindu from work_file_information where id = %s'%(file_id)

    # yuanwen = models.YuanWenTable.objects.raw(sql1)
    # yiwen = models.YiWenTable.objects.raw(sql2)
    #
    # print(yuanwen,'*'*10)
    # print(yiwen,'*'*10)
    #
    # ret['yuanwen'] = serialize('json',yuanwen)
    # ret['yiwen'] = serialize('json',yiwen)
    # print(ret)

    yuanwen_cursor = connection.cursor()
    yiwen_cursor = connection.cursor()
    num_cursor = connection.cursor()

    yuanwen_cursor.execute(sql1)
    yiwen_cursor.execute(sql2)
    num_cursor.execute(num_sql)

    yuanwens = yuanwen_cursor.fetchall()
    yiwens = yiwen_cursor.fetchall()
    nums = num_cursor.fetchall()
    ret['yuanwen'] = yuanwens
    ret['yiwens'] = yiwens
    # 进度
    ret['jindu_nums'] = nums
    print(ret)
    return JsonResponse(ret)




# 获取当前编译进度百分比
def get_progress_num(request, id):
    ret = {"code": 0}
    if request.method == 'POST':
        try:
            file = models.file_information.objects.get(pk=id)
            progress_num = file.pre_translation
            ret['data'] = progress_num
        except Exception as e:
            ret['code'] = -1
            ret['msg'] = "出现异常", e
    else:
        ret['code'] = -2
        ret['msg'] = "请求方式出错"
    return JsonResponse(ret)


# ################################### 工作台界面 ####################################### #

# API 请求原文和译文的工作台数据
def get_work_data(request):
    ret = {'code': 0}
    if request.method == 'POST':
        try:
            file_id = request.POST['file_id']
            yuanwen = models.YuanWenTable.objects.filter(file=file_id).values('pk', 'yuanwen')
            if len(yuanwen) > 1:
                yiwen = models.YiWenTable.objects.filter(file=file_id).values('pk', 'yiwen_wait_over')
                data = {
                    'yuan': serialize('json', yuanwen),
                    'yi': serialize('json', yiwen),
                }
                # data = serialize('json', data)
                ret['data'] = data
        except Exception as e:
            print("异常：", e)
            ret['code'] = -3
    else:
        ret['code'] = -4
    return JsonResponse(ret)


# ################################### 工作台操作 ####################################### #


# 获取语料库匹配多少句
def get_detail_num(request, id):
    print(id, '*'*9)
    corpus_num = models.file_information.objects.get(pk=id)
    return JsonResponse({'num': corpus_num.corpus_matching_count, 'count': corpus_num.work_count})


# 中文
# json百度翻译
@login_required
@require_http_methods(["POST"])
def bdjson(request):
    if request.method == 'POST':
        query = request.POST['query']
        if len(query) > 1:
            yiwen = fanyi(query)
        else:
            yiwen = ""
        print("翻译后：", yiwen)
        return JsonResponse({"msg": yiwen, "success": True})


# 搜狗翻译
@login_required
@require_http_methods(["POST"])
def sgjson(request):
    if request.method == 'POST':
        query = request.POST['query']
        if len(query.strip()) > 1:
            query = query.strip()
            yiwen = zh_sg(query)
        else:
            yiwen = ""
        print("翻译后：", yiwen)
        return JsonResponse({"msg": yiwen, "success": True})


# 英文
# json百度翻译
@login_required
@require_http_methods(["POST"])
def en_baidu(request):
    if request.method == 'POST':
        query = request.POST['query']
        if len(query) > 1:
            print("原"+query)
            yiwen = enfanyi(query)
            print("yi"+yiwen)
        else:
            yiwen = ""
        print("翻译后：", yiwen)
        return JsonResponse({"msg": yiwen, "success": True})


# 搜狗翻译
@login_required
@require_http_methods(["POST"])
def en_sgjson(request):
    if request.method == 'POST':
        query = request.POST['query']
        if len(query) > 1:
            query = query.strip()
            yiwen = en_sg(query)
        else:
            yiwen = ""
        print("翻译后 ：", yiwen)
        return JsonResponse({"msg": yiwen, "success": True})


# 德文
# 百度翻译
@login_required
@require_http_methods(["POST"])
def bd_de_zh(request):
    if request.method == 'POST':
        query = request.POST['query']
        if len(query) > 1:
            print("原"+query)
            yiwen = de_zh_fanyi(query)
            print("yi"+yiwen)
        else:
            yiwen = ""
        print("翻译后：", yiwen)
        return JsonResponse({"msg": yiwen, "success": True})


@login_required
@require_http_methods(["POST"])
def bd_de_en(request):
    if request.method == 'POST':
        query = request.POST['query']
        if len(query) > 1:
            print("原"+query)
            yiwen = de_en_fanyi(query)
            print("yi"+yiwen)
        else:
            yiwen = ""
        print("翻译后：", yiwen)
        return JsonResponse({"msg": yiwen, "success": True})


# 搜狗翻译
# 德-英
def sg_de_en(request):
    if request.method == 'POST':
        query = request.POST['query']
        query = query.strip()
        if len(query) > 1:
            print("原" + query)
            yiwen = de_en_sg(query)
            print("yi" + yiwen)
        else:
            yiwen = ""
        print("翻译后：", yiwen)
        return JsonResponse({"msg": yiwen, "success": True})


# 德-中
def sg_de_zh(request):
    if request.method == 'POST':
        query = request.POST['query']
        query = query.strip()
        if len(query) > 1:
            print("原" + query)
            yiwen = de_zh_sg(query)
            print("yi" + yiwen)
        else:
            yiwen = ""
        print("翻译后：", yiwen)
        return JsonResponse({"msg": yiwen, "success": True})

import json
from django.http import HttpResponse
# 单行/全部提交译文
@csrf_exempt
@login_required
@require_http_methods(["POST"])
def change_yiwen(request, id, yiwen):
    print(yiwen, id)
    # 查找译文表，把提交的译文存入到译文列和待修改列
    p = models.YiWenTable.objects.get(gl_id = id)
    p.yiwen = yiwen
    p.yiwen_wait_over = yiwen
    p.save()
    # 查找原文表，得出原文的长度
    yuanwen = models.YuanWenTable.objects.get(pk=id)
    yuanwen_len = len(yuanwen.yuanwen)
    print("原文长：", yuanwen_len)
    print(yuanwen.file.id)
    # 查找文件信息表，把提交来的译文的原文长度存到数据库中 并将是否可修改改为1（不可修改）
    file = models.file_information.objects.get(pk=yuanwen.file.id)
    print(file)
    if yuanwen.check_changeed != 1:
        yuanwen.check_changeed = 1
        file.numing += yuanwen_len
        file.paraing += 1
        print(yuanwen_len)
        print(file.numing)
        print(file.all_num)
        jindu = int((yuanwen_len+file.numing)/file.all_num*100)
        if jindu >= 100:
            jindu = 100
        file.jindu = jindu
        file.save()
        yuanwen.save()
    resp = {'numing': file.numing, 'paraing': file.paraing, 'jindu':file.jindu, "all_num":file.all_num, "all_para":file.all_para}
    return HttpResponse(json.dumps(resp), content_type="application/json")


# 全部保存译文
@login_required
@require_http_methods(["POST"])
def all_save_yiwen(request, y_id, yiwen):
    try:
        print(y_id, yiwen)
        p = models.YiWenTable.objects.get(gl_id=y_id)
        print(p)
        if request.method == 'POST':
            print(yiwen)
            p.yiwen_wait_over = yiwen
            p.save()
            return JsonResponse({"msg": "全部保存成功", "success": True})
    except:
        return JsonResponse({"msg": "全部保存失败", "success": False})


# 检查是否修改了原文栏
def yuanwen_check(request, id):
    yuanwen = models.YuanWenTable.objects.get(pk=id)
    print(yuanwen.yuanwen)
    return JsonResponse({"msg": yuanwen.yuanwen, "success": True})


# 检查是否修改了译文
def yiwen_check(request, id):
    yiwen = models.YiWenTable.objects.get(gl=id)
    print(yiwen.yiwen_wait_over)
    return JsonResponse({"msg": yiwen.yiwen_wait_over, "success": True})


# 编译原文栏后的自动机器翻译
@login_required
@require_http_methods(["POST"])
def change_content(request, query, id):
    print("oo")
    yiwen_change = fanyi(query)
    print("译文："+yiwen_change)
    yuanwen_change = query
    print("原文："+yuanwen_change)
    yiwen = models.YiWenTable.objects.get(gl_id=id)
    yuanwen = models.YuanWenTable.objects.get(pk=id)
    yiwen.yiwen_change = yiwen_change
    print('y1')
    yuanwen.yuanwen_change= yuanwen_change
    print("y2")
    yiwen.save()
    yuanwen.save()
    return JsonResponse({"msg": yiwen_change, "success": True})


# 英文 ：编译原文栏后的自动机器翻译
@login_required
@require_http_methods(["POST"])
def en_content_change(request, query, id):
    print("enoo")
    yiwen_change = enfanyi(query)
    print("译文："+yiwen_change)
    yuanwen_change = query
    print("原文："+yuanwen_change)
    yiwen = models.YiWenTable.objects.get(gl_id=id)
    yuanwen = models.YuanWenTable.objects.get(pk=id)
    yiwen.yiwen_change = yiwen_change
    print('y1')
    yuanwen.yuanwen_change= yuanwen_change
    print("y2")
    yiwen.save()
    yuanwen.save()
    return JsonResponse({"msg": yiwen_change, "success": True})


# 搜索原文
@login_required
@require_http_methods(["POST"])
def serch_yuanwen(request, file_id):
    if request.method == 'POST':
        yuanwen_keyword = request.POST['yuanwen_keyword']
        yuanwen_list = []
        yiwen_list = []
        yuanwen = models.YuanWenTable.objects.filter(file = file_id)
        for y in yuanwen:
            # print(y.yuanwen)
            if yuanwen_keyword in y.yuanwen:
                print("找到了：", y.yuanwen, y.id)
                yiwen_id = y.id
                yuanwen_list.append(y)
                yiwen = models.YiWenTable.objects.get(gl = yiwen_id)
                yiwen_list.append(yiwen)
                print(yiwen.yiwen)
        print(yuanwen_list)
        print(yiwen_list)
        wenjian = zip(yuanwen_list, yiwen_list)
        finde_count = len(yiwen_list)
        return render(request, 'work/serch.html',
                      {"wenjian": wenjian, "file_id": file_id, 'finde_count': finde_count})


# 搜索译文
@login_required
@require_http_methods(["POST"])
def serch_yiwen(request, file_id):
    if request.method == 'POST':
        yiwen_keyword = request.POST['yiwen_keyword']
        yuanwen_list = []
        yiwen_list = []
        print("接收到：", yiwen_keyword)
        yiwen = models.YiWenTable.objects.filter(file = file_id)
        for y in yiwen:
            print(y.yiwen_wait_over)
            if yiwen_keyword in y.yiwen_wait_over:
                print("找到了：", y.yiwen_wait_over, y.gl_id)
                yuanwen_id = y.gl_id
                yiwen_list.append(y)
                yuanwen = models.YuanWenTable.objects.get(pk = yuanwen_id)
                print("原文",yuanwen)
                # yuanwen = yuanwen
                yuanwen_list.append(yuanwen)
                # print(yiwen.yiwen)
        print(yuanwen_list)
        print(yiwen_list)
        wenjian = zip(yuanwen_list, yiwen_list)
        finde_count = len(yiwen_list)
        return render(request, 'work/serch.html',
                      {"wenjian": wenjian, "file_id": file_id, 'finde_count': finde_count})


# 查看修改前 修改后 前端未有展示 和入口
# 查看修改前
def change_befor(request, id):
    yiwen = models.YiWenTable.objects.get(pk=id)
    print(yiwen.yiwen)
    print(yiwen.gl.yuanwen)
    listlist = [yiwen.yiwen_wait_over, yiwen.gl.yuanwen]
    return JsonResponse({"msg": listlist, "success": True})


# 查看修改后
def change_old(request, id):
    yiwen = models.YiWenTable.objects.get(pk=id)
    print(yiwen.yiwen_change)
    print(yiwen.gl.yuanwen_change)
    listlist = [yiwen.yiwen_change, yiwen.gl.yuanwen_change]
    return JsonResponse({"msg": listlist, "success": True})


# 替换 原文
def tihuan_yuanwen(request):
    '''
        查询当前文章的所有原文 循环 替换  记录替换多少次 返回成功和替换次数
        查询当前文章的所有译文  yiwen_wait_over  循环 替换  记录替换多少次 返回成功和替换次数
    '''
    ret = {'code': 0}
    try:
        print("来了")
        f_id = request.POST['f_id']
        yuan = request.POST['yuan']
        tihuan = request.POST['tihuan']
        print(f_id, yuan, tihuan)
        print("***"*9)
        file = models.YuanWenTable.objects.filter(file_id=f_id)
        tihuan_count = 0
        for f in file:
            print(f.yuanwen)
            print("我的id：", f.id)
            print("-"*9)
            strinfo = re.compile(yuan)
            t = strinfo.sub(tihuan, f.yuanwen)
            print("我的id：", f.id)
            print("替换后：", t)
            if f.yuanwen == t:
                print('没替换')
            else:
                print("替换了")
                yuanwen = models.YuanWenTable.objects.get(pk=f.id)
                yuanwen.yuanwen = t
                yuanwen.save()
                print("原文保存OK")
                tihuan_count += 1
                change_yiwen_wait(f.id, t)
            print("######  end ######")
        print("替换了%s次"%tihuan_count)
        ret['msg'] = "成功替换了%s个记录" %tihuan_count
    except Exception as e:
        print(e)
        ret['code'] = -1
        ret['msg'] = '出现异常'
    return JsonResponse(ret)


def change_yiwen_wait(id, t):
    # 替换之后来一次机器翻译（默认百度机器翻译）
    yi_end = fanyi(t)
    yiwen_obj = models.YiWenTable.objects.get(pk=id)
    yiwen_obj.yiwen_wait_over = yi_end
    yiwen_obj.save()
    return


# 替换 译文
def tihuan_yiwen(request):
    ret = {'code': 0}
    try:
        print("来了")
        f_id = request.POST['f_id']
        yuan = request.POST['yiwen_yuan']
        tihuan = request.POST['yiwen_tihuan']
        print(f_id, yuan, tihuan)
        print("***"*9)
        file = models.YiWenTable.objects.filter(file_id=f_id)
        tihuan_count = 0
        for f in file:
            print(f.yiwen_wait_over)
            print("我的id：", f.id)
            print("-"*9)
            strinfo = re.compile(yuan)
            t = strinfo.sub(tihuan, f.yiwen_wait_over)
            print("我的id：", f.id)
            print("替换后：", t)
            if f.yiwen_wait_over == t:
                print('没替换')
            else:
                print("替换了")
                yiwen = models.YiWenTable.objects.get(pk=f.id)
                yiwen.yiwen_wait_over = t
                yiwen.save()
                tihuan_count += 1
            print("######  end ######")
        print("替换了%s次"%tihuan_count)
        ret['msg'] = "成功替换了%s个记录"%tihuan_count
    except Exception as e:
        print(e)
        ret['code'] = -1
        ret['msg'] = '出现异常'
    return JsonResponse(ret)


def default_translation(request):
    ret = {'code': 0}
    try:
        query = request.POST['query'].strip()
        print("得到要翻译的内容", query)
        default_yiwen = fanyi(query)
        print("得到的翻译：", default_yiwen)
        print(type(default_yiwen))
        ret['data'] = default_yiwen
    except Exception as e:
        print(e)
        ret['code'] = -1
        ret['msg'] = '出现异常'
    return JsonResponse(ret)


def endefault_translation(request):
    ret = {'code': 0}
    try:
        query = request.POST['query'].strip()
        print("得到要翻译的内容", query)
        default_yiwen = enfanyi(query)
        print("得到的翻译：", default_yiwen)
        print(type(default_yiwen))
        ret['data'] = default_yiwen
    except Exception as e:
        print(e)
        ret['code'] = -1
        ret['msg'] = '出现异常'
    return JsonResponse(ret)


# ################################### 导出文件 ####################################### #
# 缓冲流下载文件方法
def readFile(filename, chunk_size=512):
    with open(filename, 'rb') as f:
        while True:
            c = f.read(chunk_size)
            if c:
                yield c
            else:
                break


# 顺序读取wordneir
def read_item_block(parent):
    if isinstance(parent, _Document):
        parent_elm = parent.element.body
    elif isinstance(parent, _Cell):
        parent_elm = parent._tc
    elif isinstance(parent, _Row):
        parent_elm = parent._tr
    else:
        raise ValueError("something's not right")
    for child in parent_elm.iterchildren():
        if isinstance(child, CT_P):
            count = 1
            count_flase = 0
            res = Paragraph(child, parent)
            if res.text != '':
                yield (res,count_flase)
            else:
                try:
                    # 试着去取内联元素
                    from xml.dom.minidom import parseString
                    DOMTree = parseString(child.xml)
                    data = DOMTree.documentElement
                    nodelist = data.getElementsByTagName('pic:blipFill')
                    print('*nodelist'*9,nodelist)
                    if len(nodelist) < 1:
                        yield (res,count_flase)
                    else:
                        yield (res, count)
                except Exception as e:
                    print('*'*9,e)
                    yield (res,count_flase)
        elif isinstance(child, CT_Tbl):
            yield (Table(child, parent),)


def shuangyu(id):
    file = models.file_information.objects.get(id = id)
    file_name = file.file_name
    number = 1
    print("file_path", file.file)
    document = docx.Document(file.file)
    align = ''
    style = ''
    ccc = 0
    for block in read_item_block(document):
        if isinstance(block[0], Paragraph):
            if block[1]:
                try:
                    yiwen = models.YiWenTable.objects.get(Q(file_id=id) & Q(yiwen_number=number - 1))
                    print(yiwen.yiwen_wait_over, '*' * 9)
                    new_para = block[0].insert_paragraph_before(yiwen.yiwen_wait_over)
                    new_para.alignment = align
                    new_para.style = style
                except:
                    pass
            else:
                if ccc >= 1:
                    try:
                        yiwen = models.YiWenTable.objects.get(Q(file_id=id) & Q(yiwen_number=number-1))
                        block[0].text = re.sub(u'[^\u0020-\uD7FF\u0009\u000A\u000D\uE000-\uFFFD\U00010000-\U0010FFFF]+', '',block[0].text)
                        new_para = block[0].insert_paragraph_before(yiwen.yiwen_wait_over)
                        new_para.alignment = align
                        new_para.style = style
                    except Exception as e:
                        print(e)
                        pass
                align = block[0].alignment
                style = block[0].style
                ccc += 1
            number += 1
        elif isinstance(block[0], Table):
            for row in block[0].rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        try:
                            yiwen = models.YiWenTable.objects.get(Q(file_id=id) & Q(yiwen_number=number))
                            print(yiwen.yiwen_wait_over, '*'*9)
                            new_para = paragraph.insert_paragraph_before(paragraph.text)
                            new_para.alignment = paragraph.alignment
                            new_para.style = paragraph.style
                            paragraph.text = yiwen.yiwen_wait_over
                        except:
                            pass
                        number += 1
                ccc = 0
    print(number,id, '='*9)
    try:
        yiwen = models.YiWenTable.objects.get(Q(file_id=id) & Q(yiwen_number=number-1))
        print(yiwen.yiwen_wait_over, '*' * 9)
        new_para = document.add_paragraph(yiwen.yiwen_wait_over)
        new_para.alignment = align
        new_para.style = style
    except:
        pass

    path = "static/work/double_yiwen_doc/" + file_name
    document.save(path)
    return path


# 导出双语对照版
@login_required
def shuangyu_downloads(request, file_id):
    ret = {'code': 0}
    file = models.file_information.objects.get(pk = file_id)
    file_name = file.file_name
    print(file_name)
    file_name = '(双语对照版)' + file_name
    # 获取当前的用户
    user = request.user
    # 获取当前ID的用户
    id_user = file.file_user
    # 判断是否用户是否为一人
    if user == id_user:
        path = shuangyu(file_id)
        print('双语path', path)
        # 下载
        response = StreamingHttpResponse(readFile(path))
        response['Content-Type'] = 'application/octet-stream'
        from django.utils.http import urlquote
        file_name_chinese = file_name
        response['Content-Disposition'] = 'attachment;filename="%s"' % (urlquote(file_name_chinese))
        # response['Content-Disposition'] = 'attachment;filename="download.docx"'
        # 把状态改为已完成
        file = models.file_information.objects.get(pk=file_id)
        file.file_status = 0
        file.save()
        # ret['data'] = path
        print('导出双语对照版*************')
        return response
    else:
        ret['msg'] = "请登录自己的账号下载自己的文件"
    return JsonResponse(ret)


def word_ok(file_id):
    file = models.file_information.objects.get(pk = file_id)
    print(file.file_name)
    print(file.file)
    file_name = file.file_name
    number = 1
    print("file_path", file.file)
    document = docx.Document(file.file)
    end_num = models.YiWenTable.objects.filter(file_id=file_id).last()
    print(end_num, '-+'*9)
    print(end_num.yiwen_number, '-+'*9)
    try:
        for block in read_item_block(document):
            if isinstance(block[0], Paragraph):
                print(block)
                if block[1]:
                    pass
                else:
                    try:
                        yiwen = models.YiWenTable.objects.get(Q(file_id=file_id) & Q(yiwen_number=number))
                        print(yiwen.yiwen_wait_over)
                        block[0].text = yiwen.yiwen_wait_over
                    except:
                        pass
                number += 1
            elif isinstance(block[0], Table):
                print(block)
                for row in block[0].rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            try:
                                yiwen = models.YiWenTable.objects.get(Q(file_id=file_id) & Q(yiwen_number=number))
                                print(yiwen.yiwen_wait_over)
                                paragraph.text = yiwen.yiwen_wait_over
                            except:
                                pass
                            number += 1
            print(end_num.yiwen_number)
    except Exception as e:
        print('@#'*9, e)
        print(end_num.yiwen_number)
    print(111111111111111111)
    path = "static/work/yiwen_doc/" + file_name
    document.save(path)
    sqL_oath_save = models.file_information.objects.get(id=file_id)
    sqL_oath_save.word_path = path
    sqL_oath_save.save()
    return path, file_name


# 导出word
@login_required
# @require_http_methods(["POST"])
def downloads(request, file_id):
    ret = {'code': 0}
    file = models.file_information.objects.get(pk = file_id)
    print(file.file_name)
    # 获取当前的用户
    user = request.user
    # 获取当前ID的用户
    id_user = file.file_user
    # 判断是否用户是否为一人
    if user == id_user:
        path, file_name = word_ok(file_id)
        print(path)
        print(file_name)
        # 下载
        response = StreamingHttpResponse(readFile(path))
        response['Content-Type'] = 'application/octet-stream'
        from django.utils.http import urlquote
        file_name_chinese = file_name
        response['Content-Disposition'] = 'attachment;filename="%s"' % (urlquote(file_name_chinese))

        # 把状态改为已完成
        file = models.file_information.objects.get(pk = file_id)
        file.file_status = 0
        file.save()
        ret['data'] = path
        print('导出译文*************')
        return response
    else:
        ret['msg'] = "请登录自己的账号下载自己的文件"
        return JsonResponse(ret)


# 导出pdf---》 还未完善
# @login_required
# def downloads_pdf(request, file_id):
#     file = models.file_information.objects.get(pk = file_id)
#     print(file.file_name)
#     # 获取当前的用户
#     user = request.user
#     # 获取当前ID的用户
#     id_user = file.file_user
#     # 判断是否用户是否为一人
#     if user == id_user:
#         path, file_name = word_ok(file_id)
#         print(path)
#         print(file_name)
#         inputFile = path
        # import subprocess
        # soffice --headless --invisible --convert-to pdf /tmp/abc.docx --outdir /tmp/
        # docx_path = 'D:\code\second_Edition\\test]\问答模块.docx '
        # a = subprocess.check_output(["soffice", "--headless",
        #                              "--invisible", "--convert-to",
        #                              "docx", docx_path,
        #                              "--outdir", 'D:\code\second_Edition\\test'])
    #     outputFile = '"static/work/yiwen_doc/" + file.file_name + ".pdf"'
    #     out = to_pdf.convert_word_to_pdf(inputFile, outputFile)
    #     response = StreamingHttpResponse(readFile(out))
    #     response['Content-Type'] = 'application/octet-stream'
    #     response['Content-Disposition'] = 'attachment;filename="download.pdf"'
    #     # 把状态改为已完成
    #     file = models.file_information.objects.get(pk=file_id)
    #     file.file_status = 0
    #     file.save()
    #     return response
    # else:
    #     return render(request, 'work/update.html', {})


# ################################### 工作主页 ####################################### #
# 删除文件
@login_required
# @require_http_methods(["POST"])
def del_file(request, file_id):
    file = models.file_information.objects.get(pk = file_id)
    print(file.file_name)
    # 获取当前的用户
    user = request.user
    # 获取当前ID的用户
    id_user = file.file_user
    # 判断是否用户是否为一人
    if user == id_user:
        file = models.file_information.objects.get(pk = file_id)
        file.file_status = 2
        file.save()
        return redirect('/work/')
    else:
        return render(request, 'work/work_error.html', {})


import language_check
def yiwen_check_is_True(request):
    ret = {'code':0}
    text = request.POST['yiwen'].rstrip()
    language_type = request.POST['language_type']
    print(text, '************************************************************************************')
    print(language_type, '************************************************************************************')
    errors = []
    language_tool = language_check.LanguageTool(language_type)
    # 检测前台传过来的译文是否有错误
    matches = language_tool.check(text)

    if len(matches) > 0:
        for matche in matches:
            # 获取元素错误开始的位置
            error_index = matche.offset
            # 获取从错误那个字母开始到结束错误文本的长度
            error_length = matche.errorlength
            strs = check_texts(text,error_index,error_length)
            strs = strs.replace("\"",'')
            errors.append(strs)
            right_text = language_check.correct(text,matches)
            while '' in errors:
                errors.remove('')
            ret['data'] = errors
            ret['right_text'] = right_text
            print(ret,'-'*960)

        return JsonResponse(ret)
    else:
        ret = {'code': 1}
        return JsonResponse(ret)


def check_texts(text,index,length):
    '''

    将错误的一串文字拆分成单个字母组成一个列表，再根据索引查找其位置
    :param text: 错误的文本
    :param index: 开始的位置
    :param length: 错误字符串的长度
    :return: 返回错误的字符串
    '''
    texts = ''.join(text)
    text_list = list(texts)
    strs = ''
    while 1:
        strs += text_list[index]
        index += 1
        # text_list.remove(text_list[index])
        length -= 1
        if index > len(text_list):
            break
        if length == 0:
            break
    return strs