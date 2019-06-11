#!/usr/bin/python
# 预编译 ajax请求的 机器+语料库/// 原/译文存储数据库
from __future__ import absolute_import, unicode_literals
from work.task import enfanyi,fanyi,en_sg,zh_sg,de_en_fanyi,de_zh_fanyi,de_en_sg,de_zh_sg
from celery import task
import time
from . import views, models,corpus
import docx
from docx.text.paragraph import Paragraph
from docx.oxml.text.paragraph import CT_P
from docx.oxml.table import CT_Tbl
from docx.table import _Cell, Table, _Row
from docx import Document
from django.http import request
from django.contrib.auth.models import User
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import re

def iter_unique_cells(row):
    """Generate cells in *row* skipping empty grid cells."""
    prior_tc = None
    for cell in row.cells:
        this_tc = cell._tc
        if this_tc is prior_tc:
            continue
        prior_tc = this_tc
        yield cell


def check_all_num(text):
    '''
    检查一个文本是否是纯数字，是纯数字不翻译
    :param text:
    :return:
    '''
    reg = r'[0-9.,%]'

    t = re.findall(reg, text)
    ts = ''.join(t)

    if len(ts) == len(text):
        return 1
    else:
        return 0

# def from_trans_to(yuan, mubiao, text):
#     '''
#     选择x译x
#     :param yuan:    文件的源语言
#     :param mubiao:  文件要翻译成哪种语言
#     :param text:    待翻译的字符串
#     :return:         翻译完成的字符串
#     '''
#     if yuan == 1 and mubiao == 2:
#         # 中-英
#         yiwen = fanyi(text)
#     elif yuan == 2 and mubiao == 1:
#         # 英-中
#         yiwen = enfanyi(text)
#     elif yuan == 3 and mubiao == 1:
#         # 德-中
#         yiwen = de_zh_fanyi(text)
#     elif yuan == 3 and mubiao == 2:
#         # 德-英
#         yiwen = de_en_fanyi(text)
#     else:
#         yiwen = text
#     return yiwen
#
#
# def select_trans_type(zl, yuan, mubiao, text):
#     '''
#     选择用哪家哪种翻译
#     将yuan, mubiao, text传给from_trans_to
#     :param zl:  用户选择的翻译种类  1百度  2搜狗
#     :return:    返回接收的翻译之后的字符串
#     '''
#     if zl == 1:
#         # 百度翻译
#         yiwen = from_trans_to(yuan, mubiao, text)
#     elif zl == 2:
#         # 搜狗翻译
#         yiwen = from_trans_to(yuan, mubiao, text)
#     return yiwen
#
#
# global h_f_count
#
#
# def read_header_footer(block):
#     # 读取/翻译/保存 页眉页脚
#     global h_f_count
#     inline = block.runs
#     counts = 0
#     for i in range(len(inline)):
#         text = inline[i].text
#         # 翻译
#         yiwen = select_trans_type(zl, yuan, mubiao, text)
#         # 其中一个相对应的编号
#         read_num = h_f_count
#
#         # 存储数据库
#         # 把yuanwen_style 和yiwen_style字段 存储为是否为页眉页脚
#         line = models.YuanWenTable(yuanwen=text, yuanwen_style=1, user=user, file=word_file,
#                                    yuanwen_number=read_num)
#         line.save()
#         yw = models.YiWenTable(yiwen_wait_over=yiwen, yiwen_style=1, gl=line, user=user,
#                                file=word_file, yiwen_number=read_num)
#         yw.save()
#
#         counts += 1
#         h_f_count += 1


@task
def pre_translation(user_id,id):
    print("文件ID", id)
    paragraph_count = views.get_all_paragraph_count(id)
    # 获取到文件，文件路径 使用哪种翻译种类
    word_file = models.file_information.objects.get(pk=id)
    path = word_file.file
    zl = word_file.fanyi_zl
    # docx库 操作
    file = docx.Document(path)
    # 必须转为整形
    zl = int(zl)
    mubiao = word_file.mubiao
    yuan = word_file.yuan
    print("目标：", mubiao)
    print("翻译中->英，1==百度  2==搜狗", zl)
    print(type(zl))
    user = User.objects.get(pk=user_id)
    print("当前用户:", user)
    # 记录语料库匹配多少句
    # corpus_matching_count = 0
    read_num = 1
    # 按照顺序读取内容
    for block in views.iter_block_items(file, id, user):
        # 读取段落 并进行翻译 和存储
        if isinstance(block, Paragraph):
            # 判断是否为空
            if len(block.text) > 0:
                if not check_all_num(block.text):
                # 先对原文进行分割，然后每句分别匹配语料库，匹配不到进行机器翻译。翻译完成后 整合一段返回回来
                    if zl == 1:
                        # 百度翻译
                        # if mubiao == 1:
                        #     # 英-中
                        #     yiwen = enfanyi(block.text)
                        # elif mubiao == 2:
                        #     # 中-英
                        #     yiwen = fanyi(block.text)
                        # else:
                        #     yiwen = block.text

                        if yuan == 1 and mubiao == 2:
                            # 中-英
                            yiwen = fanyi(block.text)
                        elif yuan == 2 and mubiao == 1:
                            # 英-中
                            yiwen = enfanyi(block.text)
                        elif yuan == 3 and mubiao == 1:
                            # 德-中
                            yiwen = de_zh_fanyi(block.text)
                        elif yuan == 3 and mubiao == 2:
                            # 德-英
                            yiwen = de_en_fanyi(block.text)
                        else:
                            yiwen = block.text

                    elif zl == 2:
                        # 搜狗翻译
                        # if mubiao == 1:
                            # 英-中
                            # yiwen = en_sg(block.text)
                        # elif mubiao == 2:
                            # 中-英
                            # yiwen = zh_sg(block.text)

                        if yuan == 1 and mubiao == 2:
                            # 中-英
                            yiwen = zh_sg(block.text)
                        elif yuan == 2 and mubiao == 1:
                            # 英-中
                            yiwen = en_sg(block.text)
                        elif yuan == 3 and mubiao == 1:
                            # 德-中
                            yiwen = de_zh_sg(block.text)
                        elif yuan == 3 and mubiao == 2:
                            # 德-英
                            yiwen = de_en_sg(block.text)
                        else:
                            yiwen = block.text


                    else:
                        yiwen = block.text
                else:
                    yiwen = block.text
                print('views得到结果：', yiwen)
            else:
                yiwen = ''
            print("***** 译文 ********")
            print(yiwen)
            print("****** end  *******")
            # 一次传入多少字节
            once_num = len(block.text)
            print("一次传入多少字节", once_num)
            try:
                # 带标题的原文存入数据库
                line = models.YuanWenTable(yuanwen=block.text, yuanwen_style=block.style.name, user=user, file=word_file, yuanwen_number=read_num)
                line.save()

                # 译文存入数据库
                yw = models.YiWenTable(yiwen_wait_over=yiwen, yiwen_style=block.style.name, gl=line, user=user,
                                       file=word_file, yiwen_number=read_num)
                yw.save()
            except:
                pass
            # 将每次的字节数量+=到数据库中
            word_file.all_num += once_num
            # 传入一次即为加一段
            word_file.all_para += 1
            word_file.pre_translation = word_file.all_para / paragraph_count * 100
            print("翻译完成百分比：", word_file.pre_translation)
            word_file.save()
            read_num += 1
        # 读取表格和存储单元格数据
        elif isinstance(block, Table):
            # 循环之后有几行
            for row in block.rows:
                # 循环一行有几列
                for cell in iter_unique_cells(row):
                    # 循环某一格中的paragraphs
                    for paragraph in cell.paragraphs:
                        try:
                            style_name = paragraph.style.name
                        except:
                            style_name = 'Normal'

                        if len(paragraph.text) <= 1:
                            yiwen = ''
                        else:
                            # yiwen, count_num = corpus.split_str(paragraph.text, zl, mubiao)
                            if not check_all_num(paragraph.text):
                                if zl == 1:
                                    # 百度翻译
                                    # if mubiao == 1:
                                    #     # 英-中
                                    #     yiwen = enfanyi(block.text)
                                    # elif mubiao == 2:
                                    #     # 中-英
                                    #     yiwen = fanyi(block.text)
                                    # else:
                                    #     yiwen = block.text

                                    if yuan == 1 and mubiao == 2:
                                        # 中-英
                                        yiwen = fanyi(paragraph.text)
                                    elif yuan == 2 and mubiao == 1:
                                        # 英-中
                                        yiwen = enfanyi(paragraph.text)
                                    elif yuan == 3 and mubiao == 1:
                                        # 德-中
                                        yiwen = de_zh_fanyi(paragraph.text)
                                    elif yuan == 3 and mubiao == 2:
                                        # 德-英
                                        yiwen = de_en_fanyi(paragraph.text)
                                    else:
                                        yiwen = paragraph.text

                                elif zl == 2:
                                    # 搜狗翻译
                                    # if mubiao == 1:
                                        # 英-中
                                        # yiwen = en_sg(paragraph.text)
                                    # elif mubiao == 2:
                                        # 中-英
                                        # yiwen = zh_sg(paragraph.text)

                                    if yuan == 1 and mubiao == 2:
                                        # 中-英
                                        yiwen = zh_sg(paragraph.text)
                                    elif yuan == 2 and mubiao == 1:
                                        # 英-中
                                        yiwen = en_sg(paragraph.text)
                                    elif yuan == 3 and mubiao == 1:
                                        # 德-中
                                        yiwen = de_zh_sg(paragraph.text)
                                    elif yuan == 3 and mubiao == 2:
                                        # 德-英
                                        yiwen = de_en_sg(paragraph.text)
                                    else:
                                        yiwen = paragraph.text
                                else:
                                    yiwen = paragraph.text
                            else:
                                yiwen = paragraph.text
                            # 语料库匹配句数相加
                            # corpus_matching_count += count_num
                        line = models.YuanWenTable(yuanwen=paragraph.text, yuanwen_style=style_name, user=user, file=word_file,
                                                   yuanwen_number=read_num)
                        line.save()
                        yw = models.YiWenTable(yiwen_wait_over=yiwen, yiwen_style=style_name, gl=line, user=user,
                                               file=word_file, yiwen_number=read_num)
                        yw.save()
                        read_num += 1
            # print("语料库匹配次数：", corpus_matching_count)
            word_file.all_para += 1
            # word_file.corpus_matching_count += corpus_matching_count
            word_file.pre_translation = word_file.all_para / paragraph_count * 100
            print("翻译完成百分比：", word_file.pre_translation)
            word_file.save()
    # 正文内容读取完毕， 读取页眉页脚
    # print("运行这里")
    # for p in file.sections:
    #     print(p)
    #     for a in p.header.paragraphs:
    #         read_header_footer(a)
    #     for b in p.footer.paragraphs:
    #         read_header_footer(b)








